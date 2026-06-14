from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "UniPortal_Preliminary_Institutional_Deployment_Proposal_Sky_High_Tech.docx"
LOGO = ROOT / "client" / "public" / "branding" / "uniportal-logo.png"

BLUE = RGBColor(11, 37, 69)
BLUE_2 = RGBColor(31, 78, 121)
GOLD = RGBColor(201, 162, 39)
INK = RGBColor(33, 37, 41)
MUTED = RGBColor(93, 105, 120)
LIGHT_BLUE = "EAF1F8"
LIGHT_GOLD = "FBF5DF"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_borders(table, color="CBD5E1", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color="C9A227", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Aptos"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text, size=11, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold, italic=italic)
    return run


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 12.5, BLUE_2, 10, 5),
        ("Heading 3", 11.5, BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Aptos Display"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_footer(section):
    section.footer.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False
    for footer in (section.footer, section.first_page_footer):
        for para in footer.paragraphs:
            para.clear()
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, "Confidential - For Institutional Consultation Only | Version 1.0 - June 2026", size=8, color=MUTED)


def clear_footer(section):
    section.footer.is_linked_to_previous = False
    for para in section.footer.paragraphs:
        para.clear()
    for para in section.first_page_footer.paragraphs:
        para.clear()


def add_cover(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.55))
        p.paragraph_format.space_after = Pt(24)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(kicker, "PRELIMINARY INSTITUTIONAL PROPOSAL", size=10, color=GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(5)

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(confidential, "CONFIDENTIAL / PRELIMINARY CONSULTATION DOCUMENT", size=8.5, color=MUTED, bold=True)
    confidential.paragraph_format.space_after = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    add_text(title, "Proposal for UniPortal Institutional Digital Transformation Platform", size=25, color=BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    add_text(
        subtitle,
        "A Preliminary Institutional Proposal for University Digital Administration, Student Services & Academic Management",
        size=13,
        color=MUTED,
    )

    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="C9A227", size="18")
    rule.paragraph_format.space_after = Pt(28)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2300, 6100])
    set_borders(meta, color="D8DEE8", size="4")
    rows = [
        ("Prepared By", "Sky High Tech"),
        ("Tagline", "Technology Made Simple"),
        ("Mission Line", "Building Africa's Digital Future from South Sudan"),
        ("Document Purpose", "Institutional presentation, consultation, pilot planning, and partnership discussion"),
    ]
    for i, (label, value) in enumerate(rows):
        c0, c1 = meta.rows[i].cells
        set_cell_shading(c0, LIGHT_BLUE)
        set_cell_shading(c1, WHITE)
        p0 = c0.paragraphs[0]
        add_text(p0, label, size=9.5, color=BLUE, bold=True)
        p1 = c1.paragraphs[0]
        add_text(p1, value, size=10.5, color=INK, bold=i < 3)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(28)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(note, "Date: June 1, 2026", size=10, color=MUTED)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_before = Pt(8)
    add_text(version, "Version 1.0 - June 2026", size=8.5, color=MUTED)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(360)


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    add_text(p, text, size=16, color=BLUE, bold=True)
    paragraph_border_bottom(p, color="E3B341", size="8")
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    add_text(p, text, size=12.5, color=BLUE_2, bold=True)
    return p


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, color=BLUE)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(4)
    add_text(p, text, size=10.5)
    return p


def callout(doc, title, text, fill=LIGHT_GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9100])
    set_borders(table, color="E6D8A7", size="5")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    add_text(p, title + ": ", size=10.5, color=BLUE, bold=True)
    add_text(p, text, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def callout_bullets(doc, title, lead, items, fill=LIGHT_GOLD):
    t = doc.add_table(rows=1, cols=1)
    set_table_geometry(t, [9100])
    set_borders(t, color="E6D8A7", size="5")
    cell = t.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    add_text(p, title, size=10.5, color=BLUE, bold=True)
    p.paragraph_format.space_after = Pt(4)
    lead_p = cell.add_paragraph()
    add_text(lead_p, lead, size=10.5, color=INK)
    lead_p.paragraph_format.space_after = Pt(4)
    for item in items:
        item_p = cell.add_paragraph(style="List Bullet")
        item_p.paragraph_format.left_indent = Inches(0.2)
        item_p.paragraph_format.first_line_indent = Inches(-0.1)
        item_p.paragraph_format.space_after = Pt(2)
        add_text(item_p, item, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def table(doc, headers, rows, widths, header_fill=BLUE):
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(t, widths)
    set_borders(t, color="CCD5E1", size="5")
    hdr = t.rows[0].cells
    for i, label in enumerate(headers):
        set_cell_shading(hdr[i], "0B2545" if header_fill == BLUE else header_fill)
        p = hdr[i].paragraphs[0]
        add_text(p, label, size=9.5, color=RGBColor(255, 255, 255), bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_shading(cells[i], WHITE)
            p = cells[i].paragraphs[0]
            add_text(p, value, size=9.5, color=INK)
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return t


def implementation_roadmap(doc):
    h2(doc, "Visual Implementation Roadmap")
    t = doc.add_table(rows=3, cols=5)
    set_table_geometry(t, [2500, 450, 2500, 450, 2500])
    set_borders(t, color="B8C6D8", size="5")

    phases = [
        ("Phase 1 - Pilot", ["Registration", "Enrollment", "Student profiles", "Basic reporting"]),
        ("Phase 2 - Academic Expansion", ["Timetable", "Student support", "Academic workflows"]),
        ("Phase 3 - Institutional Expansion", ["Finance", "Examinations", "Results", "Advanced reporting", "Analytics"]),
    ]
    cols = [0, 2, 4]
    for idx, (title, items) in enumerate(phases):
        cell = t.cell(0, cols[idx])
        set_cell_shading(cell, "0B2545")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, title, size=10, color=RGBColor(255, 255, 255), bold=True)

        body_cell = t.cell(1, cols[idx])
        set_cell_shading(body_cell, LIGHT_BLUE if idx != 1 else LIGHT_GOLD)
        for p in body_cell.paragraphs:
            p.clear()
        for item in items:
            p = body_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            add_text(p, item, size=9.5, color=INK)

        foot_cell = t.cell(2, cols[idx])
        set_cell_shading(foot_cell, WHITE)
        p = foot_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer = [
            "Validate core workflow",
            "Extend academic operations",
            "Broaden institutional reporting",
        ][idx]
        add_text(p, footer, size=8.5, color=MUTED, italic=True)

    for arrow_col in [1, 3]:
        for row in range(3):
            cell = t.cell(row, arrow_col)
            set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row == 1:
                add_text(p, ">", size=16, color=GOLD, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_section_1(doc):
    h1(doc, "1. Executive Summary")
    body(doc, "University digital transformation is no longer an optional technology project. It is becoming a core institutional capability for improving student services, academic accountability, operational visibility, and leadership decision-making.")
    body(doc, "Many institutions face pressure from growing student populations, manual registration processes, fragmented records, and limited real-time reporting. These challenges affect students, registrars, academic departments, finance teams, ICT units, and senior leadership.")
    body(doc, "UniPortal responds with a focused, web-based institutional platform that demonstrates the core academic registration workflow: academic setup, student management, course enrollment, registrar review, student self-service, and foundational reporting.")
    body(doc, "This proposal positions UniPortal honestly as a current MVP with demonstrable institutional capabilities and a clear pathway toward phased deployment. Consultation is therefore essential: the system should be aligned with each university's governance model, academic structure, data policies, staff readiness, and implementation priorities.")
    callout(doc, "Executive Positioning", "UniPortal is a practical foundation for institutional digital transformation, beginning with high-value academic registration and expanding through phased consultation, pilot validation, and partnership delivery.")


def add_section_2(doc):
    h1(doc, "2. Institutional Context & Common Challenges")
    body(doc, "Across the higher education sector, many administrative processes remain partly manual, partly digital, and often distributed across offices, spreadsheets, paper files, and individual staff workflows. This context is not a criticism of any institution; it reflects common realities faced by universities working under resource, infrastructure, and growth constraints.")
    table(
        doc,
        ["Common Institutional Reality", "Operational Effect"],
        [
            ("Manual and paper-based registration", "Creates queues, repeated follow-up, and slower confirmation of student enrollment."),
            ("Delayed student enrollment processes", "Affects academic planning, course lists, and student confidence during registration periods."),
            ("Fragmented student information systems", "Makes it difficult to maintain one trusted record for students, courses, and academic sessions."),
            ("Limited academic process visibility", "Restricts leadership's ability to monitor progress, bottlenecks, and workload across offices."),
            ("Student communication challenges", "Students may not know whether action is required, approval is pending, or records are complete."),
            ("Reporting and accountability limitations", "Institutional reports can take longer to prepare and may require manual reconciliation."),
        ],
        [3300, 5800],
    )
    body(doc, "A well-managed digital platform should reduce administrative friction while respecting the institution's existing structure, policy environment, and pace of change.")


def add_section_3(doc):
    h1(doc, "3. Introduction to UniPortal")
    body(doc, "UniPortal is an institutional digital administration platform designed to support universities as they gradually modernize student services, academic management, and administrative workflows.")
    body(doc, "Its immediate purpose is to provide a controlled digital registration and enrollment environment. Its longer-term vision is to become a broader institutional platform that can support academic departments, finance workflows, examinations, results, clearance, analytics, and executive reporting.")
    table(
        doc,
        ["Dimension", "UniPortal Institutional Value"],
        [
            ("Purpose", "Digitize the core student registration and academic enrollment workflow."),
            ("Vision", "Support institutional digital transformation through phased modules."),
            ("Current MVP capability", "Demonstrates student management, course setup, enrollment, role-based access, dashboards, and reporting foundations."),
            ("Institutional value", "Improves visibility, reduces repeated manual work, and creates a structured foundation for future systems."),
            ("Scalability", "Designed for gradual expansion based on institutional priorities, readiness, and governance requirements."),
        ],
        [2100, 7000],
    )
    callout(doc, "Important Clarification", "UniPortal should not be presented as a fully completed SIS or ERP today. It is a focused institutional MVP and consultation-ready platform with a defined foundation for phased expansion.")


def add_section_4(doc):
    h1(doc, "4. Why UniPortal?")
    body(doc, "Institutions considering digital administration often face a choice between generic foreign systems, disconnected manual processes, and locally adaptable platforms. UniPortal is positioned as a practical option for institutions that need phased modernization, local support, and a system that can be shaped around institutional governance and operational realities.")
    table(
        doc,
        ["UniPortal Advantage", "Institutional Benefit"],
        [
            ("Designed with South Sudan institutional realities in mind", "Better local relevance and adaptability."),
            ("Local technical support", "Faster support, customization, and training coordination."),
            ("Phased implementation model", "Lower institutional risk through controlled pilot validation."),
            ("Consultation-driven development", "The institution shapes priorities before broad deployment."),
            ("Flexible customization", "Can align with faculties, departments, programs, roles, and institutional terminology."),
            ("Cost-aware implementation", "More realistic planning for regional institutions and staged budgets."),
        ],
        [4100, 5000],
    )
    body(doc, "The objective is institutional strengthening rather than software replacement alone. UniPortal should help improve operational efficiency, accountability, and digital service delivery while allowing the institution to modernize at a manageable pace.")


def add_section_5(doc):
    h1(doc, "5. Purpose of the Institutional Presentation & Consultation")
    body(doc, "The institutional presentation is not only a product demonstration. It is a structured consultation meeting that helps the university and Sky High Tech determine whether UniPortal fits the institution's operating realities, leadership priorities, and digital transformation roadmap.")
    h2(doc, "UniPortal Presentation Discussion Summary")
    table(
        doc,
        ["Consultation Purpose", "Expected Discussion Outcome"],
        [
            ("Demonstrate current capabilities", "Confirm what the MVP can show today through a live institutional workflow."),
            ("Gather institutional feedback", "Identify workflow gaps, policy expectations, terminology, and user requirements."),
            ("Understand university priorities", "Determine whether registration, reporting, finance, examinations, or student services should lead the roadmap."),
            ("Explore customization requirements", "Discuss faculties, departments, programs, approval rules, reports, branding, and data structures."),
            ("Discuss implementation readiness", "Review hosting, ICT capacity, data quality, staff onboarding, and pilot timing."),
            ("Align with institutional realities", "Shape a practical pilot plan that respects existing systems, staff responsibilities, and governance controls."),
        ],
        [3100, 6000],
    )
    body(doc, "This consultation-first approach reduces the risk of premature commitment and allows both parties to design a realistic pilot based on evidence rather than assumptions.")
    callout_bullets(
        doc,
        "Expected Outcome of the Presentation",
        "The institutional consultation is expected to help determine:",
        [
            "Whether UniPortal aligns with institutional priorities.",
            "Which modules should be prioritized first.",
            "Whether a pilot implementation is appropriate.",
            "Institutional focal persons.",
            "Technical readiness and hosting expectations.",
            "Preliminary next steps toward implementation planning.",
        ],
    )


def add_section_6(doc):
    h1(doc, "6. UniPortal Demonstration (Current MVP)")
    callout(doc, "Current Demonstrable MVP Features", "The features below represent capabilities available for demonstration and pilot discussion. Production deployment scope should be confirmed after institutional review, data assessment, security checks, and workflow validation.")
    items = [
        ("Student Management", ["Registration", "Student profiles", "Student records and academic identity fields"]),
        ("Course & Academic Management", ["Course creation", "Semester enrollment", "Academic setup for faculties, departments, programs, sessions, and courses"]),
        ("Role-Based Access", ["Admin access", "Registrar access", "Student access with controlled self-service views"]),
        ("Student Self-Service", ["Enrollment requests", "Course tracking", "Profile updates where institutionally permitted"]),
        ("Administrative Dashboard", ["Enrollment visibility", "Monitoring of institutional activity", "Basic reporting and operational indicators"]),
    ]
    for title, points in items:
        h2(doc, title)
        for point in points:
            bullet(doc, point)


def add_section_7(doc):
    h1(doc, "7. Institutional Vision & Future Modules")
    body(doc, "UniPortal's long-term institutional value comes from its ability to expand gradually. Future modules should be prioritized through leadership consultation, technical assessment, policy review, and pilot evidence.")
    callout(doc, "Planned Institutional Expansion Modules", "The items below are planned expansion areas, not claims that every module is already complete in the MVP.")
    table(
        doc,
        ["Expansion Area", "Institutional Purpose"],
        [
            ("Finance", "Fees, balances, payment tracking, finance reporting, and possible payment integrations."),
            ("Lecturer management", "Lecturer profiles, course assignment, grade entry workflows, and academic workload visibility."),
            ("Dean/HOD module", "Faculty and department-level oversight, approvals, and academic monitoring."),
            ("Exams & results", "Exam clearance, result processing, grade approval, transcripts, and result publication controls."),
            ("Timetable", "Course schedules, room allocation, and student/lecturer timetable visibility."),
            ("Student clearance", "Graduation, finance, library, academic, and departmental clearance workflows."),
            ("Complaint tracking", "Student support requests, escalation, resolution tracking, and service accountability."),
            ("Analytics", "Operational dashboards, trend analysis, institutional KPIs, and decision support."),
            ("Mobile accessibility", "Improved access for students and staff using phones and low-bandwidth environments."),
            ("Notifications", "Email/SMS/in-app alerts for registration status, deadlines, and institutional announcements."),
            ("AI-assisted reporting", "Future possibility for guided report generation, summaries, and leadership insights."),
        ],
        [2700, 6400],
    )
    body(doc, "Each module should be introduced only after scope, policy ownership, process design, data fields, user permissions, and support expectations are agreed.")


def add_section_8(doc):
    h1(doc, "8. Governance, Security & Institutional Data Ownership")
    callout(doc, "Data Ownership Principle", "The institution owns its data. Sky High Tech acts as a technology implementation and support partner, not the owner of institutional data.")
    table(
        doc,
        ["Governance Area", "Institutional Position"],
        [
            ("Data ownership", "All student, academic, administrative, and institutional records remain the property of the institution."),
            ("Access permissions", "Users should access only the modules and records appropriate to their official role."),
            ("Audit logging", "Critical actions should be traceable, including approvals, status changes, and administrative decisions."),
            ("Security", "Authentication, role controls, secure hosting practices, and controlled configuration should be part of deployment planning."),
            ("Cloud hosting options", "Deployment can be discussed based on institutional preference, budget, compliance, and ICT readiness."),
            ("Backup & disaster recovery", "Backup frequency, recovery process, and retention expectations should be agreed before production use."),
            ("Restricted technical access", "Technical support access should be limited, authorized, and documented according to institutional policy."),
            ("Administrative control", "The institution should retain administrative oversight of users, roles, academic structures, and data governance rules."),
        ],
        [2700, 6400],
    )
    body(doc, "Governance should be treated as a leadership and policy discussion, not only an ICT matter. Registrar, ICT, finance, academic affairs, and institutional leadership should agree on data responsibilities before rollout.")


def add_section_9(doc):
    h1(doc, "9. Institutional Deployment Approach")
    body(doc, "A phased approach is recommended so that the institution can validate workflow fit, staff readiness, data quality, user adoption, reporting needs, and technical operations before broad deployment.")
    implementation_roadmap(doc)
    table(
        doc,
        ["Phase", "Primary Scope", "Expected Institutional Outcome"],
        [
            ("Phase 1 (Pilot)", "Registration, enrollment, student profiles, and basic reporting.", "Validate the core workflow with selected faculties, departments, or student groups."),
            ("Phase 2", "Academic workflows, timetable, and student support.", "Extend operational value beyond registration and strengthen student service delivery."),
            ("Phase 3", "Finance, examinations, results, and advanced reporting.", "Move toward broader institutional administration with stronger leadership visibility."),
        ],
        [1800, 3900, 3400],
    )
    h2(doc, "Implementation Enablers")
    for item in [
        "Training for administrators, registrar staff, ICT users, and selected academic offices.",
        "Change management to help staff transition from manual or spreadsheet-based workflows.",
        "Staff onboarding with role-specific guidance and controlled user permissions.",
        "Data migration planning for students, courses, departments, programs, and academic sessions.",
        "Institutional focal persons to coordinate approvals, data validation, communication, and feedback.",
    ]:
        bullet(doc, item)


def add_section_10(doc):
    h1(doc, "10. Partnership & Sustainability Model")
    body(doc, "The recommended partnership model is collaborative and scope-based. UniPortal should be configured around the institution's academic structure, reporting expectations, governance model, and deployment readiness.")
    table(
        doc,
        ["Partnership Area", "Recommended Approach"],
        [
            ("Customization", "Adapt terminology, workflows, reports, roles, and institutional branding after consultation."),
            ("Hosting options", "Discuss cloud, managed hosting, or institution-preferred hosting based on technical and policy requirements."),
            ("Support structure", "Agree response channels, support hours, issue escalation, and responsibilities for institutional focal persons."),
            ("Maintenance", "Plan system updates, backups, monitoring, documentation, and controlled change requests."),
            ("Costing philosophy", "Keep pricing subject to institutional scope, customization requirements, hosting model, data migration, and agreed support level."),
            ("Long-term collaboration", "Use pilot evidence to define expansion modules and sustainability arrangements."),
        ],
        [2700, 6400],
    )
    callout(doc, "Pricing Position", "Final pricing should be provided only after scope confirmation. Initial discussions should focus on institutional needs, module priorities, hosting requirements, data readiness, and implementation effort.")


def add_section_11(doc):
    h1(doc, "11. Recommended Institutional Participants for Presentation")
    body(doc, "For the presentation and consultation to be effective, participation should include both decision-makers and operational users who understand the university's day-to-day administrative realities.")
    table(
        doc,
        ["Participant Group", "Reason for Attendance"],
        [
            ("University leadership", "Strategic alignment, institutional sponsorship, policy direction, and partnership decision-making."),
            ("Registrar office", "Registration workflow ownership, student records, approvals, and academic administration requirements."),
            ("ICT department", "Hosting, security, integration, support model, infrastructure, and technical readiness."),
            ("Finance office", "Future finance module requirements, payment workflows, clearance, and cost implications."),
            ("Academic affairs", "Academic calendar, course structures, enrollment rules, and quality assurance requirements."),
            ("Deans/HODs", "Faculty and department-level workflows, approvals, reporting needs, and academic oversight."),
            ("Student representatives", "Student service experience, communication needs, usability feedback, and adoption considerations."),
        ],
        [2900, 6200],
    )


def add_section_12(doc):
    h1(doc, "12. Feedback & Institutional Consultation")
    body(doc, "The consultation should produce clear institutional feedback that can be converted into a pilot scope, implementation plan, governance model, and future module roadmap.")
    table(
        doc,
        ["Consultation Question", "Response Notes"],
        [
            ("Which modules should be prioritized first?", ""),
            ("What are the institution's biggest administrative pain points?", ""),
            ("What implementation timeline is realistic for a pilot?", ""),
            ("What data governance expectations must be respected?", ""),
            ("Which existing systems, spreadsheets, or databases may need integration or migration?", ""),
            ("Which departments, faculties, or student groups should be included in the first pilot?", ""),
            ("What reports does leadership need during and after registration?", ""),
            ("Who should serve as institutional focal persons for implementation?", ""),
        ],
        [4300, 4800],
    )
    callout(doc, "Consultation Output", "After the meeting, Sky High Tech can prepare a scoped pilot note covering priority modules, deployment approach, timeline, responsibilities, support model, and cost estimate subject to agreed requirements.")


def add_section_13(doc):
    h1(doc, "13. Conclusion")
    body(doc, "UniPortal represents a practical pathway for South Sudanese universities to begin digital transformation in a focused, measurable, and institutionally responsible way. It starts with the registration and academic enrollment workflow because that workflow is visible, high-impact, and central to student service delivery.")
    body(doc, "The platform should be viewed as a long-term institutional digital transformation partner, not a one-time software installation. Its current MVP demonstrates a sufficiently defined workflow for institutional consultation, while its roadmap leaves room for finance, examinations, results, timetable, clearance, analytics, and other modules to be introduced responsibly.")
    body(doc, "Sky High Tech recommends a leadership consultation followed by a controlled pilot. This approach allows the institution to validate the system with real workflows, protect institutional data ownership, train staff, measure adoption, and plan future expansion with confidence.")
    callout(doc, "Closing Position", "UniPortal is a realistic, partnership-driven foundation for university digital administration that can grow with institutional needs.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    h1(doc, "Signature Page")
    body(doc, "Prepared by:")
    p = doc.add_paragraph()
    add_text(p, "Sky High Tech", size=13, color=BLUE, bold=True)
    body(doc, "Technology Made Simple")
    body(doc, "Building Africa's Digital Future from South Sudan")
    body(doc, "Email: skyhighttech@gmail.com")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    add_text(p, "Institution Receiving the Proposal:", size=10.5, color=BLUE, bold=True)
    p = doc.add_paragraph()
    add_text(p, "____________________________________________________________", size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_text(p, "Name & Title of Institutional Representative:", size=10.5, color=BLUE, bold=True)
    p = doc.add_paragraph()
    add_text(p, "____________________________________________________________", size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_text(p, "Date of Institutional Consultation:", size=10.5, color=BLUE, bold=True)
    p = doc.add_paragraph()
    add_text(p, "____________________________________________________________", size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    add_text(p, "Authorized Representative:", size=10.5, color=BLUE, bold=True)
    p = doc.add_paragraph()
    add_text(p, "____________________________________________________________", size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_text(p, "Date:", size=10.5, color=BLUE, bold=True)
    p = doc.add_paragraph()
    add_text(p, "____________________________________________________________", size=10.5, color=INK)


def main():
    doc = Document()
    style_doc(doc)
    add_footer(doc.sections[0])
    clear_footer(doc.sections[0])
    add_footer(doc.sections[0])
    add_cover(doc)

    for fn in [
        add_section_1,
        add_section_2,
        add_section_3,
        add_section_4,
        add_section_5,
        add_section_6,
        add_section_7,
        add_section_8,
        add_section_9,
        add_section_10,
        add_section_11,
        add_section_12,
        add_section_13,
    ]:
        fn(doc)

    for section in doc.sections:
        add_footer(section)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
