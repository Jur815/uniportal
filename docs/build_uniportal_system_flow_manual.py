from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "UniPortal_System_Flow_User_Manual.docx"
LOGO = ROOT.parent / "client" / "public" / "branding" / "uniportal-logo.png"

NAVY = "102A56"
BLUE = "1F4E79"
GOLD = "C9A227"
LIGHT_BLUE = "EAF0F8"
LIGHT_GOLD = "F8F3E3"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
INK = "172033"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=110, bottom=70, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D5DBE5", size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            cell = row.cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def set_run(run, size=9.4, bold=False, color=INK, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, before=0, after=3, line=1.08, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text, after=3, bold_lead=None):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=after)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run(rest)
    else:
        set_run(paragraph.add_run(text))
    return paragraph


def add_bullet(doc, label, detail):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph(paragraph, after=1.8, line=1.03)
    set_run(paragraph.add_run(f"{label}: "), size=9.1, bold=True, color=NAVY)
    set_run(paragraph.add_run(detail), size=9.1)
    return paragraph


def add_numbered_flow(doc, steps):
    for label, detail in steps:
        paragraph = doc.add_paragraph(style="List Number")
        set_paragraph(paragraph, after=1.8, line=1.03)
        set_run(paragraph.add_run("\t"), size=9.05)
        set_run(paragraph.add_run(f"{label}: "), size=9.05, bold=True, color=NAVY)
        set_run(paragraph.add_run(detail), size=9.05)


def style_table_text(table, header=True, body_size=8.5):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                set_paragraph(paragraph, after=0, line=1.0)
                for run in paragraph.runs:
                    set_run(
                        run,
                        size=8.35 if row_index == 0 and header else body_size,
                        bold=row_index == 0 and header,
                        color=WHITE if row_index == 0 and header else INK,
                    )


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(paragraph, before=2, after=0, line=1)
    set_run(
        paragraph.add_run(
            "Confidential Institutional Demonstration Material | UniPortal by Sky High Tech"
        ),
        size=7.6,
        color=MID_GRAY,
    )


def add_flow_diagram(doc):
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_table_geometry(table, [4200, 960, 4200])
    set_table_borders(table, color="CDD7E6", size=5)

    headers = ("FOUNDATION & ENROLLMENT", "NEXT", "RESULTS & REPORTING")
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = text
        set_cell_shading(cell, NAVY if index != 1 else GOLD)
    set_repeat_table_header(table.rows[0])

    left_steps = [
        "1  Academic Setup",
        "2  Student Registration",
        "3  Enrollment Request",
        "4  Registrar Review & Approval",
        "5  Academic Record Creation",
    ]
    right_steps = [
        "6  Lecturer Mark Submission",
        "7  Dean/HOD Review",
        "8  Registrar Final Approval & Release",
        "9  Student Result Access",
        "10  Institutional Reports",
    ]
    for row_index in range(1, 6):
        left = table.cell(row_index, 0)
        middle = table.cell(row_index, 1)
        right = table.cell(row_index, 2)
        left.text = left_steps[row_index - 1]
        right.text = right_steps[row_index - 1]
        middle.text = "↓" if row_index < 5 else "→"
        set_cell_shading(left, LIGHT_BLUE)
        set_cell_shading(right, LIGHT_GOLD)
        set_cell_shading(middle, WHITE)
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=85, start=145, bottom=85, end=145)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph(paragraph, after=0, line=1)
                for run in paragraph.runs:
                    set_run(run, size=8.8, bold=True, color=NAVY)
        middle.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in middle.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph(paragraph, after=0, line=1)
            for run in paragraph.runs:
                set_run(run, size=13, bold=True, color=GOLD)

    for index, cell in enumerate(table.rows[0].cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=75, start=100, bottom=75, end=100)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph(paragraph, after=0, line=1)
            for run in paragraph.runs:
                set_run(run, size=8, bold=True, color=NAVY if index == 1 else WHITE)


def add_two_column_scope(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_table_geometry(table, [4680, 4680])
    set_table_borders(table, color="D5DBE5", size=5)
    left, right = table.rows[0].cells
    set_cell_shading(left, LIGHT_BLUE)
    set_cell_shading(right, LIGHT_GOLD)
    set_cell_margins(left, top=120, start=150, bottom=120, end=150)
    set_cell_margins(right, top=120, start=150, bottom=120, end=150)
    set_repeat_table_header(table.rows[0])

    scope_columns = [
        (
            left,
            "CURRENT DEMONSTRATION SCOPE",
            [
                "Registration & Enrollment",
                "Academic Structure",
                "Student Management",
                "Academic Records",
                "Examination & Academic Progression",
                "Supplementary & Carry-Over Tracking",
                "Timetable",
                "Complaints / Helpdesk",
                "Exam Clearance",
                "Institutional Reports",
            ],
        ),
        (
            right,
            "PLANNED EXPANSION",
            [
                "Finance integration",
                "Payment workflows",
                "Lecturer assignment controls",
                "Advanced analytics and dashboards",
                "SMS / Email notifications",
                "Additional institutional workflows",
            ],
        ),
    ]
    for cell, heading, items in scope_columns:
        title = cell.paragraphs[0]
        set_paragraph(title, after=5, keep=True)
        set_run(title.add_run(heading), size=9.2, bold=True, color=NAVY)
        for item in items:
            paragraph = cell.add_paragraph(style="List Bullet")
            set_paragraph(paragraph, after=2.2, line=1.05)
            set_run(paragraph.add_run(item), size=8.9, color=INK)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    heading_tokens = {
        "Heading 1": (13.2, BLUE, 7, 3.5),
        "Heading 2": (10.6, NAVY, 5, 2.5),
        "Heading 3": (9.6, NAVY, 4, 2),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, left, first in (
        ("List Bullet", 0.32, -0.17),
        ("List Number", 0.32, -0.17),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(9.1)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(first)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.03


def add_title_block(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_table_geometry(table, [1500, 7860], indent_dxa=0)
    set_table_borders(table, color=WHITE, size=0)
    if LOGO.exists():
        run = table.cell(0, 0).paragraphs[0].add_run()
        run.add_picture(str(LOGO), width=Inches(0.82))
        doc_pr = run._element.xpath(".//wp:docPr")[0]
        doc_pr.set("name", "UniPortal logo")
        doc_pr.set("descr", "UniPortal institutional student information system logo")
    title_cell = table.cell(0, 1)
    title_p = title_cell.paragraphs[0]
    set_paragraph(title_p, after=1)
    set_run(
        title_p.add_run("UniPortal System Flow & User Manual"),
        size=21,
        bold=True,
        color=NAVY,
        font="Aptos Display",
    )
    subtitle = title_cell.add_paragraph()
    set_paragraph(subtitle, after=0)
    set_run(
        subtitle.add_run("Institutional Demo Quick Reference"),
        size=10.3,
        bold=True,
        color=GOLD,
    )
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=100)
    set_repeat_table_header(table.rows[0])

    rule = doc.add_paragraph()
    set_paragraph(rule, before=2, after=5)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GOLD)
    borders.append(bottom)
    p_pr.append(borders)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.56)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.26)
    configure_styles(doc)
    add_footer(section)

    add_title_block(doc)
    add_body(
        doc,
        "UniPortal is an institutional student portal and academic management platform supporting registration, enrollment, student services, academic records, examinations, progression decisions, reporting, and controlled role-based access. The current demonstration presents a practical institutional workflow that can be configured and expanded through consultation.",
        after=4,
    )

    add_heading(doc, "1. User Roles", 1)
    roles = [
        ("Admin", "Configures academic structure, courses, sessions, users, policies, dashboards, and system-wide oversight."),
        ("Registrar", "Reviews enrollment, manages records, verifies results, approves progression decisions, and releases results."),
        ("Student", "Updates permitted details, requests enrollment, follows status, uses services, and views released results."),
        ("Lecturer", "Enters course marks, reviews calculated outcomes, and submits complete semester results."),
        ("Dean/HOD", "Reviews submitted results for academic completeness and forwards them to the Registrar."),
        ("Finance (Pilot Concept)", "Demonstrates future institutional fee, financial clearance, and payment workflow possibilities. Current implementation is intentionally limited for institutional consultation and phased rollout discussions."),
    ]
    role_table = doc.add_table(rows=1, cols=2)
    role_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    role_table.rows[0].cells[0].text = "Role"
    role_table.rows[0].cells[1].text = "Primary responsibilities"
    for role, responsibility in roles:
        cells = role_table.add_row().cells
        cells[0].text = role
        cells[1].text = responsibility
        set_cell_shading(cells[0], LIGHT_BLUE)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    set_cell_shading(role_table.rows[0].cells[0], NAVY)
    set_cell_shading(role_table.rows[0].cells[1], NAVY)
    set_repeat_table_header(role_table.rows[0])
    set_fixed_table_geometry(role_table, [1650, 7710])
    set_table_borders(role_table)
    style_table_text(role_table, body_size=8.35)

    doc.add_page_break()
    add_heading(doc, "2. System Flow", 1)
    add_body(
        doc,
        "The institutional workflow links academic setup, student registration, controlled academic approvals, and leadership reporting in one traceable process.",
        after=4,
    )
    add_flow_diagram(doc)
    add_heading(doc, "How the Flow Operates", 2)
    add_numbered_flow(
        doc,
        [
            ("Academic setup", "Admin creates faculties, departments, programs, courses, sessions, and policy settings."),
            ("Student registration", "Students confirm identity and academic profiles."),
            ("Enrollment request", "Students select eligible courses and submit requests."),
            ("Registrar review", "Registrar verifies details, selection, and credits before approval or correction."),
            ("Academic records", "Approved enrollment establishes the semester academic record."),
            ("Mark submission", "Lecturer enters marks; UniPortal calculates grades, points, credits, and GPA."),
            ("Academic review", "Dean/HOD checks submitted results for academic completeness."),
            ("Final approval", "Registrar verifies standing and approval history, then releases the result."),
            ("Student access", "Students view and print only formally released results."),
            ("Reporting", "Authorized staff generate progression and department performance summaries."),
        ],
    )

    add_heading(doc, "3. Key Institutional Benefits", 1)
    benefits = [
        ("Operational efficiency", "Reduces manual paperwork and result-processing burden."),
        ("Academic accuracy", "Improves examination consistency and calculation reliability."),
        ("Progression control", "Tracks supplementary and carry-over requirements automatically."),
        ("Accountability", "Records each review, approval, change, and release decision."),
        ("Transparency", "Improves student self-service while protecting unreleased results."),
        ("Decision support", "Provides institutional reports for academic leadership."),
        ("Policy alignment", "Adapts grading, promotion, and workflow rules to the institution."),
    ]
    for label, detail in benefits:
        add_bullet(doc, label, detail)

    doc.add_page_break()
    add_heading(doc, "4. Current Demo Scope", 1)
    add_body(
        doc,
        "The current demonstration focuses on working institutional academic workflows. Planned items are presented as consultation and phased-expansion opportunities, not completed production modules.",
        after=5,
    )
    add_two_column_scope(doc)

    add_heading(doc, "5. Module Summary", 1)
    modules = [
        ("Dashboard", "Role-specific indicators, pending work, totals, and task links."),
        ("Academic Setup", "Faculties, departments, programs, sessions, and structure."),
        ("Student Management", "Identity, program placement, verification, and account status."),
        ("Course Management", "Codes, titles, credits, level, semester, links, and active status."),
        ("Enrollment Management", "Requests, Registrar decisions, reasons, tracking, and printable slips."),
        ("Academic Records", "Semester courses, grades, credits, GPA, remarks, and history."),
        ("Examinations & Progression", "Marks, calculations, standing, approvals, release, and student access."),
        ("Progression Reports", "Supplementary, carry-over, pass, repeat/discontinued, and department lists."),
        ("Timetable", "Course schedules, venues, lecturers, semesters, and programs."),
        ("Complaints / Helpdesk", "Student requests, responses, handling status, and resolution."),
        ("Exam Clearance", "Attendance, financial clearance, eligibility, notes, and approval."),
        ("Finance (Pilot Concept)", "Demonstrates future institutional fee, financial clearance, and payment workflow possibilities. Current implementation is intentionally limited for institutional consultation and phased rollout discussions."),
        ("Audit Trail & Security", "Role permissions and histories for decisions, changes, approvals, and releases."),
    ]
    module_table = doc.add_table(rows=1, cols=2)
    module_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    module_table.rows[0].cells[0].text = "Module"
    module_table.rows[0].cells[1].text = "Purpose"
    for index, (module, purpose) in enumerate(modules):
        cells = module_table.add_row().cells
        cells[0].text = module
        cells[1].text = purpose
        if index % 2 == 1:
            set_cell_shading(cells[0], LIGHT_GRAY)
            set_cell_shading(cells[1], LIGHT_GRAY)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    set_cell_shading(module_table.rows[0].cells[0], NAVY)
    set_cell_shading(module_table.rows[0].cells[1], NAVY)
    set_repeat_table_header(module_table.rows[0])
    set_fixed_table_geometry(module_table, [2760, 6600])
    set_table_borders(module_table)
    style_table_text(module_table, body_size=7.95)

    doc.add_page_break()
    add_heading(doc, "6. Examination & Academic Progression", 1)
    exam_items = [
        ("GPA / CGPA", "Calculated from course credit hours and grade points."),
        ("Grading policy", "Grade bands, pass marks, and points are institution-configurable and retained with the result."),
        ("Promotion policy", "GPA, credits, failed-course, repeat, and discontinuation thresholds follow configured rules."),
        ("Progression detection", "Failed courses are identified for supplementary or carry-over tracking; standing may be Pass, Supplementary, Carry Over, Repeat, Discontinued, or Suspended."),
        ("Approval workflow", "Lecturer Submitted → Dean/HOD Reviewed → Registrar Approved → Released."),
        ("Release control", "Students see released results only; lecturers and Dean/HOD reviewers cannot release final results."),
        ("Printable result sheet", "Shows institutional and student details, courses, marks, grades, GPA, CGPA, standing, and approval status."),
    ]
    for label, detail in exam_items:
        add_bullet(doc, label, detail)

    credentials = doc.add_table(rows=1, cols=2)
    credentials.alignment = WD_TABLE_ALIGNMENT.LEFT
    credentials.cell(0, 0).text = "DEMO CREDENTIALS"
    credentials.cell(0, 1).text = "All demo accounts use:  UniPortal123\nRole-specific emails are listed in the demo login documentation."
    set_fixed_table_geometry(credentials, [2200, 7160])
    set_table_borders(credentials, color="D8C47A", size=6)
    set_repeat_table_header(credentials.rows[0])
    set_cell_shading(credentials.cell(0, 0), GOLD)
    set_cell_shading(credentials.cell(0, 1), LIGHT_GOLD)
    style_table_text(credentials, header=False, body_size=8.8)
    for run in credentials.cell(0, 0).paragraphs[0].runs:
        set_run(run, size=8.5, bold=True, color=NAVY)
    for run in credentials.cell(0, 1).paragraphs[0].runs:
        set_run(run, size=8.8, bold=True, color=NAVY)

    add_heading(doc, "7. Institutional Configuration", 1)
    add_body(
        doc,
        "UniPortal is configurable per institution. Academic structures, grading policies, promotion rules, approval workflows, hosting, security controls, and reports can be adapted to the university's official regulations and governance requirements.",
        after=0,
    )

    doc.core_properties.title = "UniPortal System Flow & User Manual"
    doc.core_properties.subject = "Institutional demo quick reference"
    doc.core_properties.author = "Sky High Tech"
    doc.core_properties.keywords = "UniPortal, institutional demo, user manual, system flow"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
